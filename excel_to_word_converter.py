import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os

class ExcelToWordConverter:
    def __init__(self):
        self.excel_file = None
        self.sheet_names = []
        self.selected_sheet = None
        self.headers = []
        self.title_header = None
        self.sub_headers = []
        
    def upload_excel_file(self, uploaded_file):
        """Excel 파일을 업로드하고 분석합니다."""
        try:
            # Excel 파일을 메모리에서 읽기
            self.excel_file = pd.ExcelFile(uploaded_file)
            self.sheet_names = self.excel_file.sheet_names
            return True
        except Exception as e:
            st.error(f"Excel 파일을 읽는 중 오류가 발생했습니다: {str(e)}")
            return False
    
    def analyze_sheet(self, sheet_name):
        """선택된 시트의 구조를 분석합니다."""
        try:
            df = pd.read_excel(self.excel_file, sheet_name=sheet_name)
            self.selected_sheet = sheet_name
            self.headers = df.columns.tolist()
            return df
        except Exception as e:
            st.error(f"시트를 분석하는 중 오류가 발생했습니다: {str(e)}")
            return None
    
    def create_word_document(self, df, title_header, sub_headers):
        """Word 문서를 생성합니다."""
        try:
            doc = Document()
            
            # 제목 추가
            title = doc.add_heading('Excel to Word 변환 결과', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 각 행을 개별 항목으로 처리 (그룹화하지 않음)
            for idx, (_, row) in enumerate(df.iterrows()):
                # Title 헤더 (스타일 적용, 번호 X)
                title_value = row[title_header] if title_header in row else ""
                title_text = f"{title_header}: {title_value}"
                title_para = doc.add_paragraph(title_text)
                try:
                    title_para.style = doc.styles['Heading 3']
                except KeyError:
                    title_para.style = doc.styles['Heading 2']
                
                # Sub 헤더들 출력 (번호 매기기 + 값은 글머리 기호)
                numbered_index = 1
                for sub_header in sub_headers:
                    value_text = ""
                    if sub_header in row and pd.notna(row[sub_header]):
                        value_text = str(row[sub_header]).strip()
                    if not value_text:
                        continue

                    # 번호가 있는 Sub 제목
                    sub_title_para = doc.add_paragraph(f"{numbered_index}. {sub_header}")
                    sub_title_para.paragraph_format.left_indent = Inches(0.25)
                    sub_title_para_run = sub_title_para.runs[0]
                    sub_title_para_run.bold = True

                    # Sub 값은 글머리 기호 처리
                    value_para = doc.add_paragraph(value_text, style='List Bullet')
                    value_para.paragraph_format.left_indent = Inches(0.75)
                    value_para.paragraph_format.first_line_indent = Inches(-0.25)

                    numbered_index += 1

                # 각 항목 간 간격 추가 (제목 사이에만 공백 삽입)
                if idx < len(df) - 1:
                    doc.add_paragraph()
            
            return doc
        except Exception as e:
            st.error(f"Word 문서 생성 중 오류가 발생했습니다: {str(e)}")
            return None

def main():
    st.set_page_config(
        page_title="Excel to Word 변환기",
        page_icon="📄",
        layout="wide"
    )
    
    st.title("📄 Excel to Word 변환기")
    st.markdown("Excel 파일을 업로드하여 원하는 양식의 Word 문서로 변환하세요.")
    
    # 변환기 인스턴스 생성
    if 'converter' not in st.session_state:
        st.session_state.converter = ExcelToWordConverter()
    
    converter = st.session_state.converter
    
    # 1단계: Excel 파일 업로드
    st.header("1단계: Excel 파일 업로드")
    uploaded_file = st.file_uploader(
        "Excel 파일을 선택하세요",
        type=['xlsx', 'xls'],
        help=".xlsx 또는 .xls 형식의 Excel 파일을 업로드하세요."
    )
    
    if uploaded_file is not None:
        if converter.upload_excel_file(uploaded_file):
            st.success("✅ Excel 파일이 성공적으로 업로드되었습니다!")
            
            # 2단계: 시트 선택
            st.header("2단계: 변환할 시트 선택")
            selected_sheet = st.selectbox(
                "변환할 시트를 선택하세요:",
                converter.sheet_names,
                help="Word 문서로 변환할 Excel 시트를 선택하세요."
            )
            
            if selected_sheet:
                # 시트 분석
                df = converter.analyze_sheet(selected_sheet)
                if df is not None:
                    st.success(f"✅ '{selected_sheet}' 시트가 분석되었습니다!")
                    
                    # 시트 미리보기
                    st.subheader("시트 미리보기")
                    st.dataframe(df.head(10), use_container_width=True)
                    
                    # 3단계: 헤더 선택
                    st.header("3단계: 헤더 설정")
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.subheader("Title 헤더 선택")
                        st.markdown("**Title 헤더**: 번호가 매겨질 메인 헤더를 선택하세요.")
                        title_header = st.selectbox(
                            "Title 헤더:",
                            converter.headers,
                            help="이 헤더의 값들이 번호와 함께 메인 제목으로 표시됩니다."
                        )
                    
                    with col2:
                        st.subheader("Sub 헤더 선택")
                        st.markdown("**Sub 헤더**: Title 하위에 들여쓰기로 표시될 헤더들을 선택하세요.")
                        sub_headers = st.multiselect(
                            "Sub 헤더들:",
                            [h for h in converter.headers if h != title_header],
                            default=[h for h in converter.headers if h != title_header],
                            help="선택한 헤더들이 Title 하위에 들여쓰기로 표시됩니다."
                        )
                    
                    # 4단계: 변환 실행
                    if title_header and sub_headers:
                        st.header("4단계: Word 문서 생성")
                        
                        if st.button("🔄 Word 문서 생성", type="primary"):
                            with st.spinner("Word 문서를 생성하는 중..."):
                                doc = converter.create_word_document(df, title_header, sub_headers)
                                
                                if doc:
                                    # Word 문서를 바이트로 변환
                                    doc_buffer = io.BytesIO()
                                    doc.save(doc_buffer)
                                    doc_buffer.seek(0)
                                    
                                    # 다운로드 버튼
                                    st.success("✅ Word 문서가 성공적으로 생성되었습니다!")
                                    
                                    st.download_button(
                                        label="📥 Word 문서 다운로드",
                                        data=doc_buffer.getvalue(),
                                        file_name=f"converted_{selected_sheet}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                                    
                                    # 미리보기 정보
                                    st.info(f"""
                                    **생성된 문서 정보:**
                                    - 시트: {selected_sheet}
                                    - Title 헤더: {title_header}
                                    - Sub 헤더: {', '.join(sub_headers)}
                                    - 총 그룹 수: {df[title_header].nunique()}
                                    """)
                    else:
                        st.warning("⚠️ Title 헤더와 Sub 헤더를 모두 선택해주세요.")
        else:
            st.error("❌ Excel 파일 업로드에 실패했습니다.")

if __name__ == "__main__":
    main()
