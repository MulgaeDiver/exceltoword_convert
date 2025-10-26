import pandas as pd
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import threading

class ExcelToWordConverterGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Excel to Word 변환기")
        self.root.geometry("800x600")
        self.root.resizable(True, True)
        
        # 변수 초기화
        self.excel_file = None
        self.sheet_names = []
        self.selected_sheet = None
        self.headers = []
        self.df = None
        
        self.setup_ui()
        
    def setup_ui(self):
        # 메인 프레임
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 제목
        title_label = ttk.Label(main_frame, text="📄 Excel to Word 변환기", 
                               font=("Arial", 16, "bold"))
        title_label.grid(row=0, column=0, columnspan=2, pady=(0, 20))
        
        # 1단계: Excel 파일 업로드
        step1_frame = ttk.LabelFrame(main_frame, text="1단계: Excel 파일 업로드", padding="10")
        step1_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.file_path_var = tk.StringVar()
        file_entry = ttk.Entry(step1_frame, textvariable=self.file_path_var, width=60)
        file_entry.grid(row=0, column=0, padx=(0, 10), sticky=(tk.W, tk.E))
        
        browse_btn = ttk.Button(step1_frame, text="파일 선택", command=self.browse_file)
        browse_btn.grid(row=0, column=1)
        
        # 2단계: 시트 선택
        step2_frame = ttk.LabelFrame(main_frame, text="2단계: 변환할 시트 선택", padding="10")
        step2_frame.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.sheet_var = tk.StringVar()
        self.sheet_combo = ttk.Combobox(step2_frame, textvariable=self.sheet_var, width=60, state="readonly")
        self.sheet_combo.grid(row=0, column=0, sticky=(tk.W, tk.E))
        self.sheet_combo.bind('<<ComboboxSelected>>', self.on_sheet_selected)
        
        # 3단계: 헤더 설정
        step3_frame = ttk.LabelFrame(main_frame, text="3단계: 헤더 설정", padding="10")
        step3_frame.grid(row=3, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        # Title 헤더 선택
        ttk.Label(step3_frame, text="Title 헤더 (번호가 매겨질 메인 헤더):").grid(row=0, column=0, sticky=tk.W, pady=(0, 5))
        self.title_var = tk.StringVar()
        self.title_combo = ttk.Combobox(step3_frame, textvariable=self.title_var, width=60, state="readonly")
        self.title_combo.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        
        # Sub 헤더 선택
        ttk.Label(step3_frame, text="Sub 헤더들 (들여쓰기로 표시될 하위 헤더들):").grid(row=2, column=0, sticky=tk.W, pady=(0, 5))
        
        # 체크박스 프레임
        self.checkbox_frame = ttk.Frame(step3_frame)
        self.checkbox_frame.grid(row=3, column=0, sticky=(tk.W, tk.E))
        
        self.sub_vars = {}
        self.sub_checkboxes = {}
        
        # 4단계: 변환 실행
        step4_frame = ttk.LabelFrame(main_frame, text="4단계: Word 문서 생성", padding="10")
        step4_frame.grid(row=4, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))
        
        self.convert_btn = ttk.Button(step4_frame, text="🔄 Word 문서 생성", 
                                    command=self.convert_to_word, state="disabled")
        self.convert_btn.grid(row=0, column=0, pady=10)
        
        # 진행 상황 표시
        self.progress_var = tk.StringVar()
        self.progress_label = ttk.Label(step4_frame, textvariable=self.progress_var)
        self.progress_label.grid(row=1, column=0, pady=(10, 0))
        
        # 미리보기 영역
        preview_frame = ttk.LabelFrame(main_frame, text="데이터 미리보기", padding="10")
        preview_frame.grid(row=5, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(10, 0))
        
        # 트리뷰로 데이터 표시
        self.tree = ttk.Treeview(preview_frame, height=10)
        self.tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # 스크롤바
        scrollbar = ttk.Scrollbar(preview_frame, orient="vertical", command=self.tree.yview)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        self.tree.configure(yscrollcommand=scrollbar.set)
        
        # 그리드 가중치 설정
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(5, weight=1)
        step1_frame.columnconfigure(0, weight=1)
        step2_frame.columnconfigure(0, weight=1)
        step3_frame.columnconfigure(0, weight=1)
        step4_frame.columnconfigure(0, weight=1)
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(0, weight=1)
        
    def browse_file(self):
        """Excel 파일을 선택합니다."""
        file_path = filedialog.askopenfilename(
            title="Excel 파일 선택",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        
        if file_path:
            self.file_path_var.set(file_path)
            self.load_excel_file(file_path)
    
    def load_excel_file(self, file_path):
        """Excel 파일을 로드하고 분석합니다."""
        try:
            self.progress_var.set("Excel 파일을 읽는 중...")
            self.root.update()
            
            self.excel_file = pd.ExcelFile(file_path)
            self.sheet_names = self.excel_file.sheet_names
            
            # 시트 콤보박스 업데이트
            self.sheet_combo['values'] = self.sheet_names
            if self.sheet_names:
                self.sheet_combo.set(self.sheet_names[0])
                self.on_sheet_selected()
            
            self.progress_var.set("Excel 파일이 성공적으로 로드되었습니다!")
            messagebox.showinfo("성공", "Excel 파일이 성공적으로 업로드되었습니다!")
            
        except Exception as e:
            self.progress_var.set("")
            messagebox.showerror("오류", f"Excel 파일을 읽는 중 오류가 발생했습니다:\n{str(e)}")
    
    def on_sheet_selected(self, event=None):
        """시트가 선택되었을 때 호출됩니다."""
        selected_sheet = self.sheet_var.get()
        if not selected_sheet:
            return
            
        try:
            self.progress_var.set("시트를 분석하는 중...")
            self.root.update()
            
            self.df = pd.read_excel(self.excel_file, sheet_name=selected_sheet)
            self.headers = self.df.columns.tolist()
            
            # Title 헤더 콤보박스 업데이트
            self.title_combo['values'] = self.headers
            if self.headers:
                self.title_combo.set(self.headers[0])
            
            # Sub 헤더 체크박스 업데이트
            self.update_sub_headers()
            
            # 데이터 미리보기 업데이트
            self.update_preview()
            
            self.progress_var.set("시트 분석이 완료되었습니다!")
            
        except Exception as e:
            self.progress_var.set("")
            messagebox.showerror("오류", f"시트를 분석하는 중 오류가 발생했습니다:\n{str(e)}")
    
    def update_sub_headers(self):
        """Sub 헤더 체크박스들을 업데이트합니다."""
        # 기존 체크박스들 제거
        for widget in self.checkbox_frame.winfo_children():
            widget.destroy()
        
        self.sub_vars = {}
        self.sub_checkboxes = {}
        
        # 새 체크박스들 생성
        for i, header in enumerate(self.headers):
            var = tk.BooleanVar(value=True)  # 기본적으로 모두 선택
            self.sub_vars[header] = var
            
            cb = ttk.Checkbutton(self.checkbox_frame, text=header, variable=var)
            cb.grid(row=i//3, column=i%3, sticky=tk.W, padx=5, pady=2)
            self.sub_checkboxes[header] = cb
    
    def update_preview(self):
        """데이터 미리보기를 업데이트합니다."""
        if self.df is None:
            return
            
        # 기존 데이터 제거
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        # 컬럼 설정
        columns = list(self.df.columns)
        self.tree['columns'] = columns
        self.tree['show'] = 'headings'
        
        # 헤더 설정
        for col in columns:
            self.tree.heading(col, text=col)
            self.tree.column(col, width=100)
        
        # 데이터 추가 (최대 10행)
        for i, row in self.df.head(10).iterrows():
            values = [str(row[col]) if pd.notna(row[col]) else "" for col in columns]
            self.tree.insert('', 'end', values=values)
    
    def convert_to_word(self):
        """Word 문서로 변환합니다."""
        if not self.df is not None:
            messagebox.showerror("오류", "먼저 Excel 파일을 로드하고 시트를 선택해주세요.")
            return
            
        title_header = self.title_var.get()
        if not title_header:
            messagebox.showerror("오류", "Title 헤더를 선택해주세요.")
            return
            
        sub_headers = [header for header, var in self.sub_vars.items() if var.get()]
        if not sub_headers:
            messagebox.showerror("오류", "최소 하나의 Sub 헤더를 선택해주세요.")
            return
        
        # 변환을 별도 스레드에서 실행
        thread = threading.Thread(target=self._convert_worker, args=(title_header, sub_headers))
        thread.daemon = True
        thread.start()
    
    def _convert_worker(self, title_header, sub_headers):
        """변환 작업을 수행합니다."""
        try:
            self.progress_var.set("Word 문서를 생성하는 중...")
            self.convert_btn.config(state="disabled")
            self.root.update()
            
            doc = self.create_word_document(self.df, title_header, sub_headers)
            
            if doc:
                # 저장할 파일 경로 선택
                file_path = filedialog.asksaveasfilename(
                    title="Word 문서 저장",
                    defaultextension=".docx",
                    filetypes=[("Word documents", "*.docx"), ("All files", "*.*")]
                )
                
                if file_path:
                    doc.save(file_path)
                    self.progress_var.set("Word 문서가 성공적으로 생성되었습니다!")
                    messagebox.showinfo("성공", f"Word 문서가 저장되었습니다:\n{file_path}")
                else:
                    self.progress_var.set("저장이 취소되었습니다.")
            else:
                self.progress_var.set("Word 문서 생성에 실패했습니다.")
                
        except Exception as e:
            self.progress_var.set("")
            messagebox.showerror("오류", f"Word 문서 생성 중 오류가 발생했습니다:\n{str(e)}")
        finally:
            self.convert_btn.config(state="normal")
    
    def create_word_document(self, df, title_header, sub_headers):
        """Word 문서를 생성합니다."""
        try:
            doc = Document()
            
            # 제목 추가
            title = doc.add_heading('Excel to Word 변환 결과', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 각 행을 개별 항목으로 처리 (그룹화하지 않음)
            for idx, (_, row) in enumerate(df.iterrows(), 1):
                # Title 헤더 (번호 매기기)
                title_value = row[title_header] if title_header in row else ""
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(f"{idx}. {title_header}: {title_value}")
                title_run.bold = True
                
                # Sub 헤더들 출력 (들여쓰기)
                for sub_header in sub_headers:
                    if sub_header in row and pd.notna(row[sub_header]) and str(row[sub_header]).strip():
                        sub_para = doc.add_paragraph()
                        sub_para.paragraph_format.left_indent = Inches(0.5)
                        sub_run = sub_para.add_run(f"{sub_header}: {row[sub_header]}")
                        sub_run.italic = True
                
                # 각 항목 간 간격 추가
                doc.add_paragraph()
            
            return doc
        except Exception as e:
            raise e

def main():
    root = tk.Tk()
    app = ExcelToWordConverterGUI(root)
    root.mainloop()

if __name__ == "__main__":
    main()
